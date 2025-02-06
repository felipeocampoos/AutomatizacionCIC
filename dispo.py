import pandas as pd
import streamlit as st
import requests
from docx import Document
from io import StringIO, BytesIO

# Función para obtener los datos desde la API
def obtener_datos_api():
    data = {
        'token': '5A253AD70186A4E6CB1C22620A7BF6A5',
        'content': 'report',
        'format': 'csv',
        'report_id': '376',
        'csvDelimiter': ';',
        'rawOrLabel': 'label',
        'rawOrLabelHeaders': 'label',
        'exportCheckboxLabel': 'false',
        'returnFormat': 'json'
    }
    r = requests.post('https://centrodeinvestigacionesclinicas.fvl.org.co/apps/redcap/api/', data=data)
    if r.status_code == 200:
        return r.text
    else:
        st.error(f"Error al obtener los datos: {r.status_code}")
        return None

# Función para cargar y agrupar el archivo CSV
def cargar_datos(csv_data):
    df = pd.read_csv(StringIO(csv_data), sep=';')
    df_grouped = df.groupby('ID').first()
    return df_grouped

# Función para obtener la lista de coordinadores, MD asistenciales o investigadores
def obtener_coordinadores(df_grouped, columna):
    coordinadores = df_grouped[columna].dropna().unique()
    return coordinadores

# Función para calcular la disponibilidad
def calcular_disponibilidad(fase, categoria):
    disponibilidad = { 
        "Investigador Principal": {
            "Administrativo Pre inicio": "30 minutos",
            "Reclutamiento": "2 horas",
            "Reclutamiento on Hold": "1 hora",
            "Seguimiento": "2 horas",
            "Administrativo cierre": "30 minutos",
        },
        **{f"Co-Investigador {i}": {
            "Administrativo Pre inicio": "15 minutos",
            "Reclutamiento": "1 hora",
            "Reclutamiento on Hold": "30 minutos",
            "Seguimiento": "1 hora",
            "Administrativo cierre": "0 minutos",
        } for i in range(1, 8)},
        "Coordinador Principal": {
            "Administrativo Pre inicio": "1 hora",
            "Reclutamiento": "4 horas",
            "Reclutamiento on Hold": "2 horas",
            "Seguimiento": "2 horas",
            "Administrativo cierre": "1 hora",
        },
        **{f"Coordinador backup principal {i}": {
            "Administrativo Pre inicio": "0 minutos",
            "Reclutamiento": "0 minutos",
            "Reclutamiento on Hold": "0 minutos",
            "Seguimiento": "0 minutos",
            "Administrativo cierre": "0 minutos",
        } for i in range(1, 6)},
        "MD asistencial 1": {
            "Administrativo Pre inicio": "15 minutos",
            "Reclutamiento": "1 hora",
            "Reclutamiento on Hold": "30 minutos",
            "Seguimiento": "1 hora",
            "Administrativo cierre": "0 minutos",
        },
        **{f"MD asistencial {i}": {
            "Administrativo Pre inicio": "0 minutos",
            "Reclutamiento": "0 minutos",
            "Reclutamiento on Hold": "0 minutos",
            "Seguimiento": "0 minutos",
            "Administrativo cierre": "0 minutos",
        } for i in range(2, 9)},
        "Coordinador Supernumerario": {
            "Administrativo Pre inicio": "0 minutos",
            "Reclutamiento": "0 minutos",
            "Reclutamiento on Hold": "0 minutos",
            "Seguimiento": "0 minutos",
            "Administrativo cierre": "0 minutos",
        },
    }# Manteniendo la misma lógica de disponibilidad
    return disponibilidad.get(categoria, {}).get(fase, "N/A")

# Función para filtrar estudios por el coordinador, MD asistencial o investigador
def estudios_por_coordinador(df_grouped, coordinador_seleccionado, columna):
    estudios_filtrados = df_grouped[
        (df_grouped[columna] == coordinador_seleccionado) &
        (df_grouped['Estado general del estudio'].str.contains('1. Activo', na=False))
    ]
    estudios_filtrados['Estado especifico del estudio'] = estudios_filtrados['Estado especifico del estudio'].str.replace(
        r'^\d+\.\s', '', regex=True)
    
    tabla_estudios = pd.DataFrame({
        'Acrónimo': estudios_filtrados['Acrónimo Estudio'],
        'Número del Comité': estudios_filtrados['Número IRB'].apply(lambda x: f"{x:.4f}" if pd.notna(x) else ""),
        'Fase del Estudio': estudios_filtrados['Estado especifico del estudio'],
        'Sujetos Tamizados': estudios_filtrados['Total de tamizados'].fillna(0).astype(int),
        'Sujetos Activos': estudios_filtrados['Total de activos'].fillna(0).astype(int),
        'Disponibilidad de horas': estudios_filtrados.apply(
            lambda row: calcular_disponibilidad(row['Estado especifico del estudio'], columna), axis=1)
    }).reset_index(drop=True)
    
    return tabla_estudios

# Función para generar el documento en Word
def generar_documento_acumulado(reporte_acumulado):
    doc = Document()
    doc.add_heading("Reporte Acumulado", 0)

    for reporte in reporte_acumulado:
        nombre_persona = reporte['nombre']
        categoria = reporte['categoria']
        tabla_estudios = reporte['tabla']

        doc.add_heading(f"Reporte para {categoria}", level=1)
        doc.add_paragraph(f"Nombre: {nombre_persona}")

        if not tabla_estudios.empty:
            doc.add_paragraph("Tabla de estudios asociados:")
            table = doc.add_table(rows=1, cols=len(tabla_estudios.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(tabla_estudios.columns):
                hdr_cells[i].text = col_name

            for _, row in tabla_estudios.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)  # Convertir todo a string
        else:
            doc.add_paragraph("No hay estudios asociados.")

    filename = "Reporte_Acumulado.docx"
    doc.save(filename)
    return filename

# Código para ejecutar la lógica en la interfaz Streamlit
st.title("Generador de informes de disponibilidad")
st.header("Centro de Investigaciones Clínicas")
st.subheader("Fundación Valle del Lili")

if 'reporte_acumulado' not in st.session_state:
    st.session_state.reporte_acumulado = []

if st.button("Limpiar Informe"):
    st.session_state.reporte_acumulado = []
    st.success("El informe ha sido limpiado exitosamente.")

csv_data = obtener_datos_api()
if csv_data:
    df_grouped = cargar_datos(csv_data)
    
    categorias = [
        "Seleccionar", "Coordinador Principal", "Coordinador Supernumerario", "Coordinador backup principal 1", "Coordinador backup principal 2",
        "Coordinador backup principal 3", "Coordinador backup principal 4", "Coordinador backup principal 5", 
        "MD asistencial 1", "MD asistencial 2", "MD asistencial 3", "MD asistencial 4", "MD asistencial 5", 
        "MD asistencial 6", "MD asistencial 7", "MD asistencial 8", "Investigador Principal", 
        "Co-Investigador 1", "Co-Investigador 2", "Co-Investigador 3", "Co-Investigador 4", "Co-Investigador 5",
        "Co-Investigador 6", "Co-Investigador 7"
    ]  # Lista de categorías igual a la anterior
    
    seleccion_categoria = st.selectbox("Selecciona una categoría", categorias, key="categoria_general")
    
    if seleccion_categoria != "Seleccionar":
        categoria = seleccion_categoria
        coordinadores = obtener_coordinadores(df_grouped, categoria)
        
        seleccion_personas = st.multiselect("Seleccionar Personas", coordinadores, key="personas_general")
        
        if st.button("Agregar a Informe", key="agregar_general"):
            for persona in seleccion_personas:
                tabla_resultante = estudios_por_coordinador(df_grouped, persona, categoria)
                st.session_state.reporte_acumulado.append({
                    'nombre': persona,
                    'categoria': categoria,
                    'tabla': tabla_resultante
                })
            st.success(f"Se han agregado {len(seleccion_personas)} reportes al informe acumulado.")

        if st.session_state.reporte_acumulado and st.button("Generar y Descargar Informe Acumulado", key="descargar_informe"):
            filename = generar_documento_acumulado(st.session_state.reporte_acumulado)
            with open(filename, "rb") as file:
                st.download_button(label="Descargar Informe Acumulado", data=file, file_name=filename)
    else:
        st.warning("Por favor selecciona una categoría.")
    # Botón para descargar la base de datos en formato XLSX
    st.subheader("Descargar Base de Datos")
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_grouped.to_excel(writer, sheet_name='Datos', index=True)
    output.seek(0)
    
    st.download_button(
        label="Descargar Base de Datos en XLSX",
        data=output,
        file_name="Base_de_Datos.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

st.subheader("Desarrollado por: Unidad de Inteligencia Artificial- UIA")
